import requests
import io
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from standard_module import BaseModule, CommandInfo, CommandType


class DocxModule(BaseModule):
    def __init__(self):
        super().__init__('docx')
        self._first_page = True
        self._cur_document = None
        self._title = 'new_document'
        self.commands = {
            'document': CommandInfo('document', CommandType.CREATION, self._document),
            'page': CommandInfo('page', CommandType.CREATION, self._page),
            'p': CommandInfo('p', CommandType.CREATION, self._p),
            'img': CommandInfo('img', CommandType.CREATION, self._img)
        }

    def defer(self):
        self._save()

    def _document(self, properties):
        self._cur_document = Document()
        self._first_page = True
        self._title = properties['title'] if 'title' in properties else 'new_document'
        if 'orientation' in properties:
            sections = self._cur_document.sections
            if properties['orientation'] == 'landscape':
                for section in sections:
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width, section.page_height = section.page_height, section.page_width

    def _page(self, _):
        if not self._first_page:
            self._cur_document.add_page_break()
        self._first_page = False

    def _p(self, properties):
        paragraph = self._cur_document.add_paragraph()
        self._add_alignment(paragraph, properties)
        run = paragraph.add_run()
        font = run.font
        if 'size' in properties:
            font.field_size = Pt(int(properties['size']))
        if 'font' in properties:
            font.name = properties['font']
        if 'font-style' in properties:
            if properties['font-style'] == 'bold':
                run.bold = True
            elif properties['font-style'] == 'italic':
                run.italic = True
            elif properties['font-style'] == 'underline':
                run.underline = True
        text = properties['text'] if 'text' in properties else ''
        run.add_text(text)

    def _img(self, properties):
        from_web = properties['from_web'] if 'from_web' in properties else False
        image = properties['uri']
        if from_web:
            response = requests.get(properties['uri'])
            image = io.BytesIO(response.content)
        if 'format' in properties:
            from PIL import Image
            origin_img = Image.open(image)
            image = io.BytesIO()
            result_image = origin_img.convert('RGB')
            result_image.save(image, format='jpeg')
        paragraph = self._cur_document.add_paragraph()
        self._add_alignment(paragraph, properties)
        run = paragraph.add_run()
        kwargs = {}
        if 'height' in properties:
            kwargs['height'] = Inches(float(properties['height']))
        if 'width' in properties:
            kwargs['width'] = Inches(float(properties['width']))
        run.add_picture(image, **kwargs)

    def _add_alignment(self, paragraph, properties):
        paragraph_format = paragraph.paragraph_format
        if 'align' in properties:
            resolve = {
                'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            }
            paragraph_format.alignment = resolve[properties['align']]

    def _save(self):
        self._cur_document.save(self._title)
